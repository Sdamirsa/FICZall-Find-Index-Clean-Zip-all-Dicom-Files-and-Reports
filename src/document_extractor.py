#!/usr/bin/env python3
"""
document_extractor.py - Extract text from documents with OCR fallback

This module provides robust text extraction from various document formats:
- PDF files (text-based and scanned)
- DOCX files (Word documents)
- DOC files (legacy Word, limited support)
- Images (JPG, PNG, TIFF, BMP)

Features:
- Primary extraction using native libraries (fast)
- OCR fallback for scanned documents or corrupted files
- Supports English and Persian (Farsi) languages
- Configurable confidence thresholds
- Detailed extraction status reporting

Dependencies:
    pip install python-docx PyPDF2 pdf2image easyocr Pillow

For PDF to image conversion (required for OCR on PDFs):
    - Windows: Install poppler (https://github.com/osber/poppler-windows/releases)
    - macOS: brew install poppler
    - Linux: apt-get install poppler-utils

Usage:
    from src.document_extractor import DocumentExtractor, extract_text_from_file

    # Simple usage
    text, status = extract_text_from_file("document.pdf")

    # Advanced usage
    extractor = DocumentExtractor(languages=['en', 'fa'])
    result = extractor.extract("document.pdf")
    print(result['text'])
    print(result['method'])  # 'native' or 'ocr'

Author: FICZall Pipeline
Version: 1.0
"""

import os
import sys
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass, field
from enum import Enum
import tempfile
import io

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ExtractionMethod(Enum):
    """Method used for text extraction."""
    NATIVE = "native"
    OCR = "ocr"
    FAILED = "failed"
    EMPTY = "empty"


class ExtractionStatus(Enum):
    """Status of extraction attempt."""
    SUCCESS = "success"
    PARTIAL = "partial"
    OCR_FALLBACK = "ocr_fallback"
    EMPTY_DOCUMENT = "empty_document"
    CORRUPTED = "corrupted"
    UNSUPPORTED = "unsupported"
    OCR_FAILED = "ocr_failed"
    FILE_NOT_FOUND = "file_not_found"
    PERMISSION_DENIED = "permission_denied"


@dataclass
class ExtractionResult:
    """Result of text extraction."""
    text: str = ""
    status: ExtractionStatus = ExtractionStatus.SUCCESS
    method: ExtractionMethod = ExtractionMethod.NATIVE
    confidence: float = 1.0
    page_count: int = 0
    error_message: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            'text': self.text,
            'status': self.status.value,
            'method': self.method.value,
            'confidence': self.confidence,
            'page_count': self.page_count,
            'error_message': self.error_message,
            'metadata': self.metadata
        }


class DocumentExtractor:
    """
    Extract text from documents with OCR fallback.

    Supports PDF, DOCX, DOC, and image files.
    Uses EasyOCR for OCR with English and Persian support.
    """

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        'pdf': ['.pdf'],
        'docx': ['.docx'],
        'doc': ['.doc'],
        'image': ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif']
    }

    def __init__(
        self,
        languages: List[str] = None,
        use_gpu: bool = False,
        ocr_confidence_threshold: float = 0.3,
        min_text_length: int = 10,
        poppler_path: str = None
    ):
        """
        Initialize the document extractor.

        Args:
            languages: OCR languages. Default: ['en', 'fa'] (English and Persian)
            use_gpu: Use GPU for OCR (faster but requires CUDA)
            ocr_confidence_threshold: Minimum confidence for OCR results
            min_text_length: Minimum text length to consider extraction successful
            poppler_path: Path to poppler bin directory (for PDF to image conversion)
        """
        self.languages = languages or ['en', 'fa']
        self.use_gpu = use_gpu
        self.ocr_confidence_threshold = ocr_confidence_threshold
        self.min_text_length = min_text_length
        self.poppler_path = poppler_path

        # Lazy loading for OCR reader
        self._ocr_reader = None
        self._ocr_available = None

        # Check available libraries
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which dependencies are available."""
        self.has_docx = False
        self.has_pypdf2 = False
        self.has_pdfplumber = False
        self.has_pdf2image = False
        self.has_easyocr = False
        self.has_pillow = False

        try:
            from docx import Document
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX support disabled.")

        try:
            import PyPDF2
            self.has_pypdf2 = True
        except ImportError:
            logger.warning("PyPDF2 not installed. PDF text extraction may be limited.")

        try:
            import pdfplumber
            self.has_pdfplumber = True
        except ImportError:
            pass  # Optional, PyPDF2 is sufficient

        try:
            from pdf2image import convert_from_path
            self.has_pdf2image = True
        except ImportError:
            logger.warning("pdf2image not installed. PDF OCR fallback disabled.")

        try:
            import easyocr
            self.has_easyocr = True
        except ImportError:
            logger.warning("easyocr not installed. OCR support disabled.")

        try:
            from PIL import Image
            self.has_pillow = True
        except ImportError:
            logger.warning("Pillow not installed. Image processing disabled.")

    @property
    def ocr_reader(self):
        """Lazy load OCR reader."""
        if self._ocr_reader is None and self.has_easyocr:
            try:
                import easyocr
                logger.info(f"Initializing EasyOCR with languages: {self.languages}")
                self._ocr_reader = easyocr.Reader(
                    self.languages,
                    gpu=self.use_gpu,
                    verbose=False
                )
                self._ocr_available = True
                logger.info("EasyOCR initialized successfully")
            except Exception as e:
                logger.error(f"Failed to initialize EasyOCR: {e}")
                self._ocr_available = False
        return self._ocr_reader

    @property
    def ocr_available(self) -> bool:
        """Check if OCR is available."""
        if self._ocr_available is None:
            _ = self.ocr_reader  # Trigger lazy load
        return self._ocr_available or False

    def get_file_type(self, file_path: Union[str, Path]) -> Optional[str]:
        """Get file type from extension."""
        ext = Path(file_path).suffix.lower()
        for file_type, extensions in self.SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                return file_type
        return None

    def extract(self, file_path: Union[str, Path]) -> ExtractionResult:
        """
        Extract text from a document file.

        Args:
            file_path: Path to the document file

        Returns:
            ExtractionResult with text and metadata
        """
        file_path = Path(file_path)

        # Check if file exists
        if not file_path.exists():
            return ExtractionResult(
                status=ExtractionStatus.FILE_NOT_FOUND,
                method=ExtractionMethod.FAILED,
                error_message=f"File not found: {file_path}"
            )

        # Check file type
        file_type = self.get_file_type(file_path)
        if file_type is None:
            return ExtractionResult(
                status=ExtractionStatus.UNSUPPORTED,
                method=ExtractionMethod.FAILED,
                error_message=f"Unsupported file type: {file_path.suffix}"
            )

        # Extract based on file type
        try:
            if file_type == 'pdf':
                return self._extract_pdf(file_path)
            elif file_type == 'docx':
                return self._extract_docx(file_path)
            elif file_type == 'doc':
                return self._extract_doc(file_path)
            elif file_type == 'image':
                return self._extract_image(file_path)
        except PermissionError as e:
            return ExtractionResult(
                status=ExtractionStatus.PERMISSION_DENIED,
                method=ExtractionMethod.FAILED,
                error_message=str(e)
            )
        except Exception as e:
            logger.error(f"Extraction error for {file_path}: {e}")
            return ExtractionResult(
                status=ExtractionStatus.CORRUPTED,
                method=ExtractionMethod.FAILED,
                error_message=str(e)
            )

    def _extract_pdf(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF file."""
        text = ""
        page_count = 0

        # Try native extraction first
        if self.has_pdfplumber:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                logger.warning(f"pdfplumber failed for {file_path}: {e}")

        if not text and self.has_pypdf2:
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    page_count = len(reader.pages)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                logger.warning(f"PyPDF2 failed for {file_path}: {e}")

        # Check if we got meaningful text
        text = text.strip()
        if len(text) >= self.min_text_length:
            return ExtractionResult(
                text=text,
                status=ExtractionStatus.SUCCESS,
                method=ExtractionMethod.NATIVE,
                page_count=page_count
            )

        # Fallback to OCR
        logger.info(f"Native extraction failed or empty, trying OCR for {file_path}")
        return self._extract_pdf_ocr(file_path, page_count)

    def _extract_pdf_ocr(self, file_path: Path, page_count: int = 0) -> ExtractionResult:
        """Extract text from PDF using OCR."""
        if not self.has_pdf2image:
            return ExtractionResult(
                status=ExtractionStatus.OCR_FAILED,
                method=ExtractionMethod.FAILED,
                error_message="pdf2image not installed. Cannot OCR PDF.",
                page_count=page_count
            )

        if not self.ocr_available:
            return ExtractionResult(
                status=ExtractionStatus.OCR_FAILED,
                method=ExtractionMethod.FAILED,
                error_message="OCR not available. Install easyocr.",
                page_count=page_count
            )

        try:
            from pdf2image import convert_from_path

            # Convert PDF to images
            convert_kwargs = {}
            if self.poppler_path:
                convert_kwargs['poppler_path'] = self.poppler_path

            images = convert_from_path(file_path, **convert_kwargs)
            page_count = len(images)

            all_text = []
            total_confidence = 0
            result_count = 0

            for i, image in enumerate(images):
                logger.info(f"OCR processing page {i+1}/{page_count}")

                # Convert PIL image to numpy array for easyocr
                import numpy as np
                image_np = np.array(image)

                # Run OCR
                results = self.ocr_reader.readtext(image_np)

                page_text = []
                for (bbox, text, confidence) in results:
                    if confidence >= self.ocr_confidence_threshold:
                        page_text.append(text)
                        total_confidence += confidence
                        result_count += 1

                if page_text:
                    all_text.append(' '.join(page_text))

            text = '\n\n'.join(all_text).strip()
            avg_confidence = total_confidence / result_count if result_count > 0 else 0

            if len(text) >= self.min_text_length:
                return ExtractionResult(
                    text=text,
                    status=ExtractionStatus.OCR_FALLBACK,
                    method=ExtractionMethod.OCR,
                    confidence=avg_confidence,
                    page_count=page_count
                )
            else:
                return ExtractionResult(
                    text=text,
                    status=ExtractionStatus.EMPTY_DOCUMENT,
                    method=ExtractionMethod.OCR,
                    confidence=avg_confidence,
                    page_count=page_count,
                    error_message="OCR extracted very little text"
                )

        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            return ExtractionResult(
                status=ExtractionStatus.OCR_FAILED,
                method=ExtractionMethod.FAILED,
                error_message=str(e),
                page_count=page_count
            )

    def _extract_docx(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX file."""
        if not self.has_docx:
            return ExtractionResult(
                status=ExtractionStatus.UNSUPPORTED,
                method=ExtractionMethod.FAILED,
                error_message="python-docx not installed"
            )

        try:
            from docx import Document
            doc = Document(file_path)

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            text = '\n'.join(paragraphs).strip()

            if len(text) >= self.min_text_length:
                return ExtractionResult(
                    text=text,
                    status=ExtractionStatus.SUCCESS,
                    method=ExtractionMethod.NATIVE,
                    page_count=1  # DOCX doesn't have clear page boundaries
                )
            else:
                return ExtractionResult(
                    text=text,
                    status=ExtractionStatus.EMPTY_DOCUMENT,
                    method=ExtractionMethod.NATIVE,
                    error_message="Document contains very little text"
                )

        except Exception as e:
            logger.warning(f"DOCX extraction failed for {file_path}: {e}")
            return ExtractionResult(
                status=ExtractionStatus.CORRUPTED,
                method=ExtractionMethod.FAILED,
                error_message=str(e)
            )

    def _extract_doc(self, file_path: Path) -> ExtractionResult:
        """Extract text from legacy DOC file."""
        # Try using antiword or textract if available
        try:
            import subprocess

            # Try antiword (Linux/Mac)
            try:
                result = subprocess.run(
                    ['antiword', str(file_path)],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    text = result.stdout.strip()
                    return ExtractionResult(
                        text=text,
                        status=ExtractionStatus.SUCCESS,
                        method=ExtractionMethod.NATIVE
                    )
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

            # Try catdoc (Linux/Mac)
            try:
                result = subprocess.run(
                    ['catdoc', str(file_path)],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    text = result.stdout.strip()
                    return ExtractionResult(
                        text=text,
                        status=ExtractionStatus.SUCCESS,
                        method=ExtractionMethod.NATIVE
                    )
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

        except Exception as e:
            logger.warning(f"DOC extraction failed for {file_path}: {e}")

        return ExtractionResult(
            status=ExtractionStatus.UNSUPPORTED,
            method=ExtractionMethod.FAILED,
            error_message="Legacy .doc format requires antiword or catdoc. Convert to .docx."
        )

    def _extract_image(self, file_path: Path) -> ExtractionResult:
        """Extract text from image using OCR."""
        if not self.ocr_available:
            return ExtractionResult(
                status=ExtractionStatus.OCR_FAILED,
                method=ExtractionMethod.FAILED,
                error_message="OCR not available. Install easyocr."
            )

        if not self.has_pillow:
            return ExtractionResult(
                status=ExtractionStatus.OCR_FAILED,
                method=ExtractionMethod.FAILED,
                error_message="Pillow not installed"
            )

        try:
            from PIL import Image
            import numpy as np

            # Load image
            image = Image.open(file_path)

            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')

            image_np = np.array(image)

            # Run OCR
            results = self.ocr_reader.readtext(image_np)

            texts = []
            total_confidence = 0
            result_count = 0

            for (bbox, text, confidence) in results:
                if confidence >= self.ocr_confidence_threshold:
                    texts.append(text)
                    total_confidence += confidence
                    result_count += 1

            text = ' '.join(texts).strip()
            avg_confidence = total_confidence / result_count if result_count > 0 else 0

            if len(text) >= self.min_text_length:
                return ExtractionResult(
                    text=text,
                    status=ExtractionStatus.SUCCESS,
                    method=ExtractionMethod.OCR,
                    confidence=avg_confidence,
                    page_count=1
                )
            else:
                return ExtractionResult(
                    text=text,
                    status=ExtractionStatus.EMPTY_DOCUMENT,
                    method=ExtractionMethod.OCR,
                    confidence=avg_confidence,
                    page_count=1,
                    error_message="OCR extracted very little text"
                )

        except Exception as e:
            logger.error(f"Image OCR failed for {file_path}: {e}")
            return ExtractionResult(
                status=ExtractionStatus.OCR_FAILED,
                method=ExtractionMethod.FAILED,
                error_message=str(e)
            )


def extract_text_from_file(
    file_path: Union[str, Path],
    languages: List[str] = None,
    use_gpu: bool = False
) -> Tuple[str, str]:
    """
    Simple function to extract text from a file.

    Args:
        file_path: Path to the document
        languages: OCR languages (default: ['en', 'fa'])
        use_gpu: Use GPU for OCR

    Returns:
        Tuple of (extracted_text, status_message)
    """
    extractor = DocumentExtractor(languages=languages, use_gpu=use_gpu)
    result = extractor.extract(file_path)

    status_msg = f"{result.status.value}"
    if result.method == ExtractionMethod.OCR:
        status_msg += f" (OCR, confidence: {result.confidence:.2f})"
    if result.error_message:
        status_msg += f" - {result.error_message}"

    return result.text, status_msg


# =============================================================================
# CLI Interface
# =============================================================================

def main():
    """Command-line interface for document extraction."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Extract text from documents with OCR fallback',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python document_extractor.py document.pdf
    python document_extractor.py document.docx --output output.txt
    python document_extractor.py image.jpg --languages en fa
    python document_extractor.py scanned.pdf --gpu
        """
    )

    parser.add_argument(
        'file',
        type=str,
        help='Path to document file'
    )

    parser.add_argument(
        '--output', '-o',
        type=str,
        help='Output file path (default: print to stdout)'
    )

    parser.add_argument(
        '--languages', '-l',
        nargs='+',
        default=['en', 'fa'],
        help='OCR languages (default: en fa)'
    )

    parser.add_argument(
        '--gpu',
        action='store_true',
        help='Use GPU for OCR'
    )

    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Verbose output'
    )

    parser.add_argument(
        '--json',
        action='store_true',
        help='Output result as JSON'
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Extract text
    extractor = DocumentExtractor(
        languages=args.languages,
        use_gpu=args.gpu
    )

    result = extractor.extract(args.file)

    if args.json:
        import json
        output = json.dumps(result.to_dict(), ensure_ascii=False, indent=2)
    else:
        output = result.text
        if args.verbose:
            print(f"Status: {result.status.value}", file=sys.stderr)
            print(f"Method: {result.method.value}", file=sys.stderr)
            if result.confidence < 1.0:
                print(f"Confidence: {result.confidence:.2f}", file=sys.stderr)
            if result.error_message:
                print(f"Error: {result.error_message}", file=sys.stderr)
            print("---", file=sys.stderr)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output)
        print(f"Output saved to: {args.output}")
    else:
        print(output)

    # Return exit code based on status
    if result.status in [ExtractionStatus.SUCCESS, ExtractionStatus.OCR_FALLBACK]:
        return 0
    elif result.status == ExtractionStatus.EMPTY_DOCUMENT:
        return 0  # Not an error, just empty
    else:
        return 1


if __name__ == '__main__':
    sys.exit(main())
