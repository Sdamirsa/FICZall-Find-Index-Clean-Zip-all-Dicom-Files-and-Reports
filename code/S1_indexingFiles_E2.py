import os
import json
import logging
import tempfile
import asyncio
import argparse
import multiprocessing
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Dict, List, Any, Optional
import time
from datetime import datetime

# 3rd-party libs
try:
    import pydicom
except ImportError:
    logger.warning("pydicom not installed. DICOM processing will be unavailable.")
    pydicom = None

try:
    from docx import Document as DocxDocument
except ImportError:
    logger.warning("python-docx not installed. DOCX processing will be unavailable.")
    DocxDocument = None

try:
    import PyPDF2
except ImportError:
    logger.warning("PyPDF2 not installed. PDF processing will be unavailable.")
    PyPDF2 = None

from tqdm import tqdm

###############################################################################
# LOGGING CONFIGURATION
###############################################################################
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s"
)
logger = logging.getLogger(__name__)

###############################################################################
# ENUM-LIKE STATUS
###############################################################################
class FolderStatus:
    PENDING = "pending"
    FILES_DONE = "files_done"
    DICOM_DONE = "dicom_done"
    ERROR = "error"

###############################################################################
# HELPER FUNCTIONS
###############################################################################
def atomic_write_json(data: Dict, path: str, pretty: bool = True) -> None:
    dir_name = os.path.dirname(path)
    if dir_name.strip():
        os.makedirs(dir_name, exist_ok=True)

    with tempfile.NamedTemporaryFile(dir=dir_name, suffix=".tmp", delete=False, mode='w', encoding="utf-8") as tmp_file:
        temp_path = tmp_file.name
        try:
            if pretty:
                json.dump(data, tmp_file, indent=4, ensure_ascii=False)
            else:
                json.dump(data, tmp_file, ensure_ascii=False)
            tmp_file.flush()
            tmp_file.close()  # Ensure file is closed before moving
            os.replace(temp_path, path)
        except Exception as e:
            logger.error(f"Error during atomic_write_json for {path}: {e}")
            try:
                tmp_file.close()
            except:
                pass
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
            raise

def read_json(path: str) -> Dict:
    if not os.path.isfile(path):
        raise ValueError(f"JSON file {path} not found.")
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Error reading JSON {path}: {e}")
        return {}

def make_folder_key(root_path: Path, folder_path: Path) -> str:
    rel = folder_path.relative_to(root_path)
    return str(rel).replace(os.sep, "_._")

def sanitize_filename(filename: str, max_length: int = 255) -> str:
    illegal_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|', '\0']
    for char in illegal_chars:
        filename = filename.replace(char, '_')
    sanitized = ''.join('_' if ord(char) < 32 else char for char in filename)
    if len(sanitized) > max_length and '.' in sanitized:
        name, ext = sanitized.rsplit('.', 1)
        sanitized = f"{name[:max_length - len(ext) - 1]}.{ext}"
    else:
        sanitized = sanitized[:max_length]
    return sanitized.strip(' .') or "unnamed_file"


###############################################################################
# FOLDER AND FILE DISCOVERY
###############################################################################
def build_all_folders_json(root_dir: str, output_json: str, overwrite: bool = False) -> None:
    if not overwrite and os.path.exists(output_json):
        logger.info(f"Skipping build_all_folders_json; using existing file: {output_json}")
        return
    root_path = Path(root_dir).resolve()
    logger.info(f"Building folder structure from {root_path}")
    folders_dict = {
        make_folder_key(root_path, Path(dirpath)): {
            "path": dirpath,
            "status": FolderStatus.PENDING,
            "error": ""
        }
        for dirpath, _, _ in tqdm(os.walk(root_path), desc="Scanning folders")
    }
    atomic_write_json(folders_dict, output_json)

def build_all_files_json(folders_json: str, output_json: str, write_interval: int = 30) -> None:
    all_folders = read_json(folders_json)
    if not all_folders:
        logger.warning(f"No folders found in {folders_json}. Nothing to process.")
        return

    # Initialize empty dict if file doesn't exist
    if not os.path.exists(output_json):
        all_files = {}
    else:
        all_files = read_json(output_json)
    
    folder_keys = list(all_folders.keys())
    last_write = 0

    for i, folder_key in enumerate(tqdm(folder_keys, desc="Building file records")):
        folder_info = all_folders[folder_key]
        folder_path = folder_info["path"]

        if folder_info["status"] != FolderStatus.PENDING:
            continue

        if folder_key in all_files and all_files[folder_key]["status"] != FolderStatus.PENDING:
            continue

        all_files[folder_key] = {
            "path": folder_path,
            "status": FolderStatus.FILES_DONE,
            "error": folder_info.get("error", ""),
            "dicom_files": [],
            "document_files": [],
            "image_files": []
        }

        try:
            with os.scandir(folder_path) as entries:
                for entry in entries:
                    if entry.is_file():
                        ext = entry.name.lower().split('.')[-1] if '.' in entry.name else ''
                        file_path = str(Path(entry.path).resolve())
                        if ext == "dcm" or ext == "":
                            all_files[folder_key]["dicom_files"].append(file_path)
                        elif ext in ["pdf", "doc", "docx"]:
                            all_files[folder_key]["document_files"].append(file_path)
                        elif ext in ["png", "jpg", "jpeg"]:
                            all_files[folder_key]["image_files"].append(file_path)
        except Exception as e:
            logger.error(f"Error reading folder '{folder_path}': {e}")
            all_files[folder_key]["status"] = FolderStatus.ERROR
            all_files[folder_key]["error"] += f"\nError reading folder '{folder_path}': {e}"

    atomic_write_json(all_files, output_json)

###############################################################################
# DICOM FILE PROCESSING
###############################################################################

def read_dicom_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from a DICOM file with error handling."""
    if not pydicom:
        raise ImportError("pydicom is required for DICOM processing")
    
    try:
        dataset = pydicom.dcmread(file_path, stop_before_pixels=True, force=False)
    except Exception as e:
        logger.warning(f"Failed to read DICOM file {file_path}: {e}")
        # Try with force=True as fallback
        try:
            dataset = pydicom.dcmread(file_path, stop_before_pixels=True, force=True)
        except Exception as e2:
            logger.error(f"Failed to read DICOM file even with force=True {file_path}: {e2}")
            raise

    # Extract Pixel Spacing (0028,0030) - can be a list or single value
    pixel_spacing = getattr(dataset, "PixelSpacing", None)
    if pixel_spacing:
        if hasattr(pixel_spacing, '__iter__') and not isinstance(pixel_spacing, str):
            # Convert list/array to string format (e.g., "1.0\1.0")
            pixel_spacing_str = "\\".join(str(x) for x in pixel_spacing)
        else:
            pixel_spacing_str = str(pixel_spacing)
    else:
        pixel_spacing_str = "Unknown"
    
    # Extract Image Type (0008,0008) - can be a multivalue field
    image_type = getattr(dataset, "ImageType", None)
    if image_type:
        if hasattr(image_type, '__iter__') and not isinstance(image_type, str):
            # Convert list/array to string format (e.g., "ORIGINAL\PRIMARY\LOCALIZER\CT_SOM5 TOP")
            image_type_str = "\\".join(str(x) for x in image_type)
        else:
            image_type_str = str(image_type)
    else:
        image_type_str = "Unknown"

    return {
        "file_path": file_path,
        "patient_name": str(getattr(dataset, "PatientName", "Unknown")),
        "patient_id": str(getattr(dataset, "PatientID", "Unknown")),
        "patient_age": str(getattr(dataset, "PatientAge", "Unknown")),
        "study_date": str(getattr(dataset, "StudyDate", "Unknown")),
        "study_time": str(getattr(dataset, "StudyTime", "Unknown")),
        "study_description": str(getattr(dataset, "StudyDescription", "Unknown")),  # NEW: (0008,1030)
        "study_modality": str(getattr(dataset, "Modality", "Unknown")),
        "study_body_part": str(getattr(dataset, "BodyPartExamined", "Unknown")),
        "study_series_number": str(getattr(dataset, "SeriesNumber", "Unknown")),
        "study_series_description": str(getattr(dataset, "SeriesDescription", "Unknown")),
        "study_protocol": str(getattr(dataset, "ProtocolName", "Unknown")),
        "station_name": str(getattr(dataset, "StationName", "Unknown")),
        "institution_name": str(getattr(dataset, "InstitutionName", "Unknown")),
        "study_comments": str(getattr(dataset, "StudyComments", "Unknown")),
        "series_number": str(getattr(dataset, "SeriesNumber", "Unknown")),
        "slice_thickness": str(getattr(dataset, "SliceThickness", "Unknown")),
        "image_orientation_patient": str(getattr(dataset, "ImageOrientationPatient", "Unknown")),
        "image_position_patient": str(getattr(dataset, "ImagePositionPatient", "Unknown")),
        # NEW FIELDS REQUESTED:
        "pixel_spacing": pixel_spacing_str,  # (0028,0030) Pixel Spacing
        "image_type": image_type_str,        # (0008,0008) Image Type
        "slice_location": str(getattr(dataset, "SliceLocation", "Unknown"))  # (0020,1041) Slice Location
    }

async def process_dicom_files(all_files_json: str, concurrency: int = 3, batch_size: int = 50, output_directory: str = ".") -> None:
    """Process DICOM files with improved batching and error handling."""
    if not pydicom:
        logger.error("pydicom not available. Skipping DICOM processing.")
        return
        
    all_files = read_json(all_files_json)
    if not all_files:
        return

    loop = asyncio.get_event_loop()
    executor = ThreadPoolExecutor(max_workers=concurrency)
    processed_count = 0
    error_count = 0
    folder_keys = list(all_files.keys())

    async def process_one_folder(folder_key: str) -> None:
        folder_info = all_files[folder_key]
        if folder_info["status"] == FolderStatus.DICOM_DONE:
            return

        dicom_files = folder_info.get("dicom_files", [])
        if not dicom_files:
            folder_info["status"] = FolderStatus.DICOM_DONE
            return

        # Initialize counters
        processed_count = 0
        error_count = 0
        study_dict = {}
        for i in range(0, len(dicom_files), batch_size):
            batch = dicom_files[i:i+batch_size]
            tasks = [loop.run_in_executor(executor, read_dicom_metadata, path) for path in batch]
            results = await asyncio.gather(*tasks, return_exceptions=True)

            for result in results:
                if isinstance(result, Exception):
                    error_count += 1
                    folder_info["error"] += f"\nDICOM error: {result}"
                    continue
                processed_count += 1
                key = (result["study_date"], result["patient_id"], result["station_name"] or result["institution_name"])
                study_dict.setdefault(key, []).append(result)

        for (sd, pid, center), dicoms in study_dict.items():
            raw_name = f"{sd}_._{pid}_._{center or 'Unknown'}_._Study.jsonl"
            filename = os.path.join(output_directory,sanitize_filename(raw_name))
            
            with open(filename, 'a', encoding='utf-8') as f:
                f.writelines(json.dumps(d, ensure_ascii=False) + '\n' for d in dicoms)

        if folder_info["status"] != FolderStatus.ERROR:
            folder_info["status"] = FolderStatus.DICOM_DONE

    # Process folders with progress tracking
    with tqdm(total=len(folder_keys), desc="Processing DICOM metadata") as pbar:
        tasks = [process_one_folder(folder_key) for folder_key in folder_keys]
        for coro in asyncio.as_completed(tasks):
            await coro
            pbar.update(1)

    atomic_write_json(all_files, all_files_json)
    logger.info(f"DICOM processing complete: {processed_count} files processed, {error_count} errors")

###############################################################################
# DOCUMENT TEXT EXTRACTION
###############################################################################
def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF with improved error handling."""
    if not PyPDF2:
        return "PyPDF2 not installed"
    
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            if reader.is_encrypted:
                logger.warning(f"PDF is encrypted: {file_path}")
                return ""
            text_parts = []
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {i} from {file_path}: {e}")
            return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction error for {file_path}: {e}")
        return ""

def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX with improved error handling."""
    if not DocxDocument:
        return "python-docx not installed"
    
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        # Extract from paragraphs
        for par in doc.paragraphs:
            if par.text.strip():
                text_parts.append(par.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error for {file_path}: {e}")
        return ""

def extract_doc_text(file_path: str) -> str:
    return "Extraction for .doc is not implemented."

async def extract_documents_text(all_files_json: str, output_json: str, concurrency: int = 3) -> None:
    all_files = read_json(all_files_json)
    # Initialize documents_text as empty dict if file doesn't exist
    if not os.path.exists(output_json):
        documents_text = {}
    else:
        documents_text = read_json(output_json)
    
    loop = asyncio.get_event_loop()
    executor = ThreadPoolExecutor(max_workers=concurrency)

    for folder_key in tqdm(all_files.keys(), desc="Extracting document text"):
        folder_info = all_files[folder_key]
        if folder_info["status"] == FolderStatus.ERROR:
            continue

        doc_files = folder_info.get("document_files", [])
        tasks = []
        paths = []
        for path in doc_files:
            if path in documents_text:
                continue
            ext = path.lower().split('.')[-1]
            if ext == "pdf":
                tasks.append(loop.run_in_executor(executor, extract_pdf_text, path))
            elif ext == "docx":
                tasks.append(loop.run_in_executor(executor, extract_docx_text, path))
            elif ext == "doc":
                tasks.append(loop.run_in_executor(executor, extract_doc_text, path))
            paths.append(path)

        if not tasks:
            continue

        results = await asyncio.gather(*tasks, return_exceptions=True)
        for path, result in zip(paths, results):
            if isinstance(result, Exception):
                folder_info["status"] = FolderStatus.ERROR
                folder_info["error"] += f"\nDoc error: {result}"
            else:
                documents_text[path] = result

    atomic_write_json(documents_text, output_json)
    atomic_write_json(all_files, all_files_json)
    
###############################################################################
# PIPELINE ENTRY POINT
###############################################################################
def S1_run_pipeline(root_directory: str, output_directory: str = ".", overwrite: bool = False):
    """Main entry point for the S1 pipeline."""
    # Check for required dependencies
    missing_deps = []
    if not pydicom:
        missing_deps.append("pydicom")
    if not PyPDF2:
        missing_deps.append("PyPDF2")
    if not DocxDocument:
        missing_deps.append("python-docx")
    
    if missing_deps:
        logger.warning(f"Missing optional dependencies: {', '.join(missing_deps)}")
        logger.warning("Some functionality may be limited. Install with: pip install " + ' '.join(missing_deps))
    
    asyncio.run(main_workflow(root_directory, output_directory, overwrite))

async def main_workflow(root_dir: str, output_dir: str, overwrite: bool = False):
    cpu_count = multiprocessing.cpu_count()
    dicom_conc = max(1, min(cpu_count - 1, 8))
    print(f"using {dicom_conc} cpu for dicom_conc")
    doc_conc = max(1, min(cpu_count - 1, 4))
    print(f"using {doc_conc} cpu for doc_conc")

    build_all_folders_json(root_dir, os.path.join(output_dir, "S1_indexingFiles_allFolders.json"), overwrite=overwrite)
    build_all_files_json(os.path.join(output_dir, "S1_indexingFiles_allFolders.json"), os.path.join(output_dir, "S1_indexingFiles_allFiles.json"), write_interval=20)
    await process_dicom_files(os.path.join(output_dir, "S1_indexingFiles_allFiles.json"), concurrency=dicom_conc, output_directory=output_dir)
    await extract_documents_text(os.path.join(output_dir, "S1_indexingFiles_allFiles.json"), os.path.join(output_dir, "S1_indexingFiles_allDocuments.json"), concurrency=doc_conc)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Index and extract metadata from DICOM files and documents",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("--root", help="Root directory to scan")
    parser.add_argument("--output_dir", help="Output directory for JSON files")
    parser.add_argument("--overwrite", action="store_true", help="Whether to overwrite existing JSONs")
    parser.add_argument("--concurrency", type=int, help="Override number of concurrent workers")
    args = parser.parse_args()

    # Get root directory if not provided
    root_dir = args.root
    if not root_dir:
        root_dir = input("Enter root directory to scan: ")
        while not os.path.isdir(root_dir):
            print(f"Directory '{root_dir}' does not exist.")
            root_dir = input("Enter valid root directory to scan: ")
    
    # Get output directory if not provided
    output_dir = args.output_dir
    if not output_dir:
        output_dir = input("Enter output directory (or press Enter for current directory): ") or "."
        os.makedirs(output_dir, exist_ok=True)
    
    S1_run_pipeline(root_dir, output_directory=output_dir, overwrite=args.overwrite)